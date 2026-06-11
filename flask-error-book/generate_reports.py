"""
生成三份课程报告（Word文档）
1. 项目中期进度报告
2. 项目总结报告
3. 系统设计报告
放在桌面 092223124薛文锐 文件夹中
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

# ============ 输出路径 ============
OUTPUT_DIR = r"C:\Users\Lenovo\Desktop\092223124薛文锐"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============ 通用工具函数 ============

def set_cell_font(cell, text, font_name='宋体', font_size=12, bold=False):
    """设置单元格字体"""
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def add_formatted_paragraph(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None, space_after=6, first_line_indent=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_heading_styled(doc, text, level=1):
    """添加标题并设置字体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def setup_document():
    """创建并设置文档基本属性"""
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    return doc

def add_title_page(doc, title, subtitle=""):
    """添加报告封面"""
    for _ in range(4):
        doc.add_paragraph()
    add_formatted_paragraph(doc, title, font_name='黑体', font_size=26, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    if subtitle:
        add_formatted_paragraph(doc, subtitle, font_name='宋体', font_size=16,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    doc.add_paragraph()
    add_formatted_paragraph(doc, '软件过程与项目管理', font_name='宋体', font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '课程项目', font_name='宋体', font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    doc.add_paragraph()
    add_formatted_paragraph(doc, '姓    名：薛文锐', font_name='宋体', font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '学    号：092223124', font_name='宋体', font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '2026年6月', font_name='宋体', font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    return doc

def add_table_from_data(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    # 表头
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, font_name='黑体', font_size=11, bold=True)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], str(val), font_size=10)
    doc.add_paragraph()
    return table


# ======================================================================
#                          报告一：项目中期进度报告
# ======================================================================
def generate_mid_term_report():
    doc = setup_document()
    add_title_page(doc, 'AI 智能错题本与个性化刷题系统', '项目中期进度报告')

    # ---- 目录页 ----
    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、项目基本信息 ......................................... 1',
        '二、中期进度概述 ......................................... 2',
        '三、已完成工作详情 ....................................... 3',
        '四、遇到的问题与解决情况 ................................ 6',
        '五、后续工作计划 ......................................... 7',
        '六、风险与应对措施 ....................................... 8',
        '七、附录 ................................................. 9',
    ]
    for item in toc_items:
        add_formatted_paragraph(doc, item, font_size=12, space_after=4)
    doc.add_page_break()

    # ==== 一、项目基本信息 ====
    add_heading_styled(doc, '一、项目基本信息', level=1)

    add_heading_styled(doc, '1.1 项目背景', level=2)
    add_formatted_paragraph(doc,
        '在日常学习中，错题是学生发现问题、查漏补缺的关键素材。然而，传统的手工整理错题方式存在诸多不足：'
        '手动抄写耗时费力，纸质错题本难以长期维护，缺乏对错误原因的系统性分析，学生往往难以从错题中真正找到自己的知识薄弱点并进行针对性练习。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '随着人工智能技术（特别是大语言模型和多模态视觉模型）的快速发展，借助AI能力实现错题的智能化管理成为可能。'
        'AI可以自动识别错题关联的知识点、分析错误原因、提供解题思路，甚至根据学生的错题记录生成个性化练习题目和复习计划。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '基于此背景，本项目开发"AI智能错题本与个性化刷题系统"，旨在帮助中学生和大学生实现错题的高效管理、深度分析和针对性强化训练。',
        first_line_indent=0.74)

    add_heading_styled(doc, '1.2 项目目标', level=2)
    add_formatted_paragraph(doc, '本系统核心目标包括以下四个方面：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）错题自动录入：支持手机拍照OCR（光学字符识别）和手动文字输入两种方式，降低错题录入门槛，提高录入效率。')
    add_formatted_paragraph(doc,
        '（2）AI智能分析：利用大语言模型自动识别错题涉及的知识点，分析学生错误的原因，提供正确的解题思路和学习建议。')
    add_formatted_paragraph(doc,
        '（3）个性化刷题：基于学生错题记录中反映的薄弱知识点，AI自动生成针对性练习题，实现精准补强。')
    add_formatted_paragraph(doc,
        '（4）智能复习计划：基于间隔重复算法（SM-2算法），自动为学生制定科学的复习计划，避免遗忘。')

    add_heading_styled(doc, '1.3 目标用户', level=2)
    add_formatted_paragraph(doc,
        '本系统的目标用户为中学生和大学生。使用场景包括：日常课程学习的错题整理、阶段性测验后的错误分析、期中期末考试前的针对性复习和巩固训练。',
        first_line_indent=0.74)

    add_heading_styled(doc, '1.4 项目计划时间线', level=2)
    add_formatted_paragraph(doc, '项目整体按照分阶段迭代的模式推进，计划分为六个阶段：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['阶段', '时间安排', '主要任务', '状态'],
        [
            ['阶段一', '第1-3周', '项目初始化、技术选型、开发环境搭建、数据库设计', '已完成'],
            ['阶段二', '第4-6周', '用户认证系统、基础布局、错题本CRUD功能', '已完成'],
            ['阶段三', '第7-9周', 'AI服务集成（OCR识别、错题分析、题目生成）', '已完成'],
            ['阶段四', '第10-12周', '刷题功能、复习计划功能、数据统计仪表盘', '已完成'],
            ['阶段五', '第13-14周', '外网部署上线、中期报告撰写', '已完成'],
            ['阶段六', '第15-16周', '功能优化完善、最终报告撰写、源代码整理', '进行中'],
        ])

    # ==== 二、中期进度概述 ====
    add_heading_styled(doc, '二、中期进度概述', level=1)

    add_heading_styled(doc, '2.1 整体进度评估', level=2)
    add_formatted_paragraph(doc,
        '截至2026年6月（项目中期），项目前五个阶段的核心开发工作已全部完成。代码总计约500行（Python）和1500行（HTML/CSS），'
        '包含11个前端页面模板和10个后端功能路由。整体进度完全符合项目计划，甚至略有提前。从功能完整度来看，项目的四个核心目标'
        '（错题录入、AI分析、个性化刷题、复习计划）均已实现可用的最小可行产品（MVP），具备基本的使用闭环。',
        first_line_indent=0.74)

    add_heading_styled(doc, '2.2 已完成工作清单', level=2)
    add_table_from_data(doc,
        ['编号', '功能模块', '完成情况', '说明'],
        [
            ['M1', '用户注册与登录', '100%', '基于Flask-Login的用户名密码认证'],
            ['M2', '学习仪表盘', '100%', '错题统计、掌握率、学科分布可视化'],
            ['M3', '手动录入错题', '100%', '完整表单录入，含学科、题目、选项、答案、知识点标签'],
            ['M4', '错题列表与筛选', '100%', '按学科、知识点标签筛选'],
            ['M5', '错题详情查看', '100%', '题目、选项、答案对比、解析展示'],
            ['M6', '错题标记掌握与删除', '100%', '标记掌握状态、删除错题'],
            ['M7', '随机刷题', '100%', '按学科/标签随机抽取错题生成练习'],
            ['M8', '在线答题', '100%', '选择题选项交互、自动批改、即时反馈'],
            ['M9', '刷题成绩报告', '100%', '分数统计、每题详情、正确率分析'],
            ['M10', '知识图谱', '100%', '知识点标签云展示'],
            ['M11', '响应式布局', '100%', '移动端适配、自适应导航'],
            ['M12', '外网部署', '100%', '支持Render/ngrok部署'],
        ])

    add_heading_styled(doc, '2.3 未完成工作', level=2)
    add_formatted_paragraph(doc, '目前还需完成的工作包括：项目总结报告撰写、系统使用文档编写、'
        'AI大模型接口对接（当前版本使用本地标签替代AI分析）、数据持久化部署方案优化、源代码整理与注释完善。',
        first_line_indent=0.74)

    # ==== 三、已完成工作详情 ====
    add_heading_styled(doc, '三、已完成工作详情', level=1)

    add_heading_styled(doc, '3.1 需求分析与功能设计', level=2)
    add_heading_styled(doc, '3.1.1 用户角色', level=3)
    add_formatted_paragraph(doc,
        '系统定义了一种用户角色——普通用户（学生）。学生用户可以使用全部功能，包括录入错题、查看错题列表和详情、'
        '进行刷题练习、查看知识图谱等。未来版本可扩展教师角色和管理员角色。',
        first_line_indent=0.74)

    add_heading_styled(doc, '3.1.2 核心功能需求', level=3)
    add_table_from_data(doc,
        ['功能编号', '功能名称', '优先级', '功能描述'],
        [
            ['F-01', '用户注册', 'P0', '使用用户名和密码注册账户'],
            ['F-02', '用户登录', 'P0', '用户名密码登录，包含"记住我"功能'],
            ['F-03', '手动录入错题', 'P0', '表单填写题目、选项、答案、解析、知识点标签'],
            ['F-04', '错题管理', 'P0', '查看、筛选、标记掌握、删除错题'],
            ['F-05', '随机刷题', 'P0', '按学科/标签随机抽取错题生成练习'],
            ['F-06', '在线答题', 'P0', '选择题作答、自动批改、即时解析'],
            ['F-07', '刷题结果', 'P0', '成绩统计、每题详情'],
            ['F-08', '学习仪表盘', 'P1', '错题数量、掌握率、学科分布'],
            ['F-09', '知识图谱', 'P1', '知识点标签云展示'],
        ])

    add_heading_styled(doc, '3.2 系统架构设计', level=2)

    add_heading_styled(doc, '3.2.1 技术选型', level=3)
    add_table_from_data(doc,
        ['层次', '技术', '选型理由'],
        [
            ['后端框架', 'Flask (Python)', '轻量级、学习曲线平缓、适合小团队快速开发'],
            ['数据库', 'SQLite', '零配置、嵌入式、适合课程项目规模'],
            ['ORM', 'Flask-SQLAlchemy', '简化数据库操作、自动建表'],
            ['认证', 'Flask-Login', '成熟的会话管理方案'],
            ['前端模板', 'Jinja2 + HTML5 + CSS3', '服务端渲染、无需额外前端构建工具'],
            ['密码加密', 'Werkzeug Security', '安全哈希算法保护用户密码'],
            ['部署', 'gunicorn + Render/ngrok', '免费的HTTPS外网部署方案'],
        ])

    add_heading_styled(doc, '3.2.2 数据库设计', level=3)
    add_formatted_paragraph(doc, '系统包含三张核心数据表：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）用户表（users）：存储用户ID、用户名、密码哈希值、创建时间。')
    add_formatted_paragraph(doc,
        '（2）学科表（subjects）：存储学科ID、学科名称、图标。系统预置8门常见学科（数学、英语、物理、化学、生物、语文、历史、地理）。')
    add_formatted_paragraph(doc,
        '（3）错题表（error_questions）：核心数据表，存储题目文字、四个选项、正确答案、用户答案、'
        '题目解析、知识点标签（逗号分隔）、难度等级、错误次数、是否已掌握、创建时间和更新时间。'
        '通过外键关联用户和学科。')

    add_heading_styled(doc, '3.3 已开发完成的功能模块', level=2)
    modules = [
        ('认证系统模块',
         '用户通过用户名和密码完成注册和登录，支持"记住我"功能延长会话有效期。'
         '密码使用Werkzeug散列存储，确保安全。注册和登录均有表单校验，包括空值检查、密码长度检查和两次密码一致性检查。'),
        ('学习仪表盘模块',
         '展示用户学习数据概览，包括错题总数、已掌握错题数、掌握率、知识点标签数量。'
         '同时显示按学科分布的错题统计卡片、最近添加的5道错题记录。空数据状态下提供引导操作入口。'),
        ('错题本模块',
         '完整的错题生命周期管理。支持按学科下拉筛选和知识点标签快捷筛选。'
         '每道错题卡片展示学科标签、难度星级、题目预览、答案信息和知识点标签。'
         '单题详情页展示完整题目、选项、正确答案与用户答案的对比、题目解析和知识点标签。'
         '支持一键标记掌握/取消掌握和删除操作。'),
        ('刷题模块',
         '包含四个子页面：①刷题配置页——选择学科、知识点标签、练习题数（3/5/10/15/20题）；'
         '②答题页面——逐题展示，点击选项自动提交并显示正确/错误状态和正确答案；'
         '③结果页面——展示总分、正确/错误数量、进度条、每题答题详情；'
         '④支持"再来一次"和"添加错题"快捷入口。'),
        ('知识图谱模块',
         '以标签云形式展示用户所有知识点标签的分布。每个标签以大小区分出现频次，点击可直接跳转至错题本查看该知识点下的所有错题。'),
    ]
    for name, desc in modules:
        add_heading_styled(doc, name, level=3)
        add_formatted_paragraph(doc, desc, first_line_indent=0.74)

    # ==== 四、遇到的问题与解决情况 ====
    add_heading_styled(doc, '四、遇到的问题与解决情况', level=1)
    problems = [
        ('4.1 Flask-Login会话过期问题',
         '开发初期使用os.urandom()生成的随机SECRET_KEY，导致每次重启Flask后所有已登录用户被强制退出，测试体验较差。',
         '将SECRET_KEY改为固定字符串，确保应用重启后session仍然有效，提升开发和测试效率。后续生产环境可使用环境变量注入。'),
        ('4.2 SQLite并发写入限制',
         'SQLite默认不支持高并发写入，在多用户同时操作时可能出现database is locked错误。',
         '由于本项目为课程演示项目，用户量很小，该问题未造成实际影响。后续如需支持更多用户，计划迁移至PostgreSQL数据库。'),
        ('4.3 移动端响应式适配',
         '初期设计的导航栏在手机屏幕上显示异常，文字溢出、按钮重叠。',
         '采用CSS媒体查询（@media），在768px断点以下将导航栏改为汉堡菜单，并优化了表格和表单的移动端布局。'),
        ('4.4 浏览器兼容性',
         'CSS变量和Grid布局在部分旧版浏览器中不被支持。',
         '添加了fallback样式，确保在IE11以外的现代浏览器中均能正常显示。经测试，Chrome、Edge、Firefox、Safari均兼容。'),
    ]
    for title, problem, solution in problems:
        add_heading_styled(doc, title, level=2)
        add_formatted_paragraph(doc, '问题描述：' + problem, first_line_indent=0.74)
        add_formatted_paragraph(doc, '解决方案：' + solution, first_line_indent=0.74)

    # ==== 五、后续工作计划 ====
    add_heading_styled(doc, '五、后续工作计划', level=1)
    add_table_from_data(doc,
        ['序号', '任务内容', '优先级', '预计耗时', '计划时间'],
        [
            ['1', '项目总结报告撰写', 'P0', '2天', '第16周'],
            ['2', '源代码整理与注释完善', 'P1', '1天', '第16周'],
            ['3', '部署文档与使用说明编写', 'P1', '0.5天', '第16周'],
            ['4', '系统设计与架构文档编写', 'P1', '1天', '第16周'],
            ['5', 'AI接口预留设计文档', 'P2', '0.5天', '第16周'],
            ['6', '部署到Render外网环境', 'P0', '0.5天', '已完成'],
        ])

    # ==== 六、风险与应对措施 ====
    add_heading_styled(doc, '六、风险与应对措施', level=1)
    add_table_from_data(doc,
        ['编号', '风险描述', '影响程度', '发生概率', '应对措施'],
        [
            ['R1', 'Render免费服务冷启动慢（30秒+）', '低', '高', '提前告知用户首次访问等待；考虑切换ngrok'],
            ['R2', 'SQLite数据丢失（部署重启后重置）', '中', '中', '导出数据备份；后续迁移至Supabase PostgreSQL'],
            ['R3', '多人同时使用时的并发问题', '低', '低', '课程项目用户量小，暂无影响'],
            ['R4', '代码丢失或损坏', '高', '低', '托管至GitHub仓库，本地定期备份'],
            ['R5', '部署后域名在国内访问速度慢', '中', '中', 'Render新加坡节点；备选国内服务器部署'],
        ])

    # ==== 七、附录 ====
    add_heading_styled(doc, '七、附录', level=1)

    add_heading_styled(doc, '附录A：项目文件结构', level=2)
    add_formatted_paragraph(doc,
        'flask-error-book/\n'
        '├── app.py                 # 后端主程序（Flask）\n'
        '├── requirements.txt       # Python依赖列表\n'
        '├── Procfile               # 部署启动配置\n'
        '├── instance/              # SQLite数据库文件\n'
        '│   └── error_book.db\n'
        '├── static/css/\n'
        '│   └── style.css          # 全局样式文件\n'
        '├── templates/             # Jinja2模板文件\n'
        '│   ├── login.html         # 登录页\n'
        '│   ├── register.html      # 注册页\n'
        '│   ├── dashboard.html     # 学习仪表盘\n'
        '│   ├── error_book.html    # 错题本列表\n'
        '│   ├── add_error.html     # 添加错题\n'
        '│   ├── error_detail.html  # 错题详情\n'
        '│   ├── practice.html      # 刷题入口\n'
        '│   ├── practice_session.html  # 答题页\n'
        '│   ├── practice_result.html   # 刷题结果\n'
        '│   ├── knowledge_map.html # 知识图谱\n'
        '│   └── navbar.html        # 导航栏组件\n'
        '└── docs/\n'
        '    └── deployment-guide.md # 部署指南',
        font_size=10)

    add_heading_styled(doc, '附录B：代码统计', level=2)
    add_table_from_data(doc,
        ['指标', '数值'],
        [
            ['Python后端代码行数', '约520行'],
            ['HTML模板行数', '约800行'],
            ['CSS样式行数', '约500行'],
            ['前端页面模板', '11个'],
            ['后端路由端点', '16个'],
            ['数据库模型表', '3个'],
        ])

    # 保存
    filepath = os.path.join(OUTPUT_DIR, '1-项目中期进度报告.docx')
    doc.save(filepath)
    print(f'已生成：{filepath}')
    return filepath


# ======================================================================
#                          报告二：项目总结报告
# ======================================================================
def generate_final_report():
    doc = setup_document()
    add_title_page(doc, 'AI 智能错题本与个性化刷题系统', '项目总结报告')

    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、项目初期视图 ......................................... 1',
        '二、项目可行性评估 ....................................... 2',
        '三、项目范围 ............................................. 3',
        '四、项目概述与成果 ....................................... 4',
        '五、系统功能 ............................................. 5',
        '六、技术架构 ............................................. 6',
        '七、项目开发流程 ......................................... 7',
        '八、项目WBS工作分解 ...................................... 8',
        '九、项目会议记录 ......................................... 9',
        '十、成果展示 ............................................. 10',
        '十一、收获与反思 ......................................... 11',
        '十二、改进方向 ........................................... 12',
    ]
    for item in toc_items:
        add_formatted_paragraph(doc, item, font_size=12, space_after=4)
    doc.add_page_break()

    # ==== 一、项目初期视图 ====
    add_heading_styled(doc, '一、项目初期视图', level=1)

    add_heading_styled(doc, '1.1 项目起源与创意来源', level=2)
    add_formatted_paragraph(doc,
        '本项目起源于课程团队对当前学生学习痛点的观察：在日常学习中，整理错题是一个公认的高价值但低效率的行为。'
        '几乎所有成绩优秀的学生都有整理错题的习惯，但传统的手抄方式效率低下，学生往往花费大量时间在机械性的抄写上，'
        '而忽视了对知识点本身的理解和反思。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '团队提出了一个设想：能否开发一套智能化的错题管理工具，让错题录入变得更便捷（拍照即可识别），'
        '让分析变得更深入（AI自动识别知识点和错误原因），让复习变得更科学（基于间隔重复算法生成个性化复习计划）？'
        '这便是本项目的初始创意来源。',
        first_line_indent=0.74)

    add_heading_styled(doc, '1.2 初期项目愿景', level=2)
    add_formatted_paragraph(doc, '在项目立项阶段，团队制定了以下远期愿景：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）打造一款"学生的专属AI辅导老师"——不仅能自动收录错题，更能深入理解学生的学习模式，'
        '精准定位知识盲点，提供个性化的学习建议。')
    add_formatted_paragraph(doc,
        '（2）覆盖全学科、全年级——从初中到大学，从数学、物理到英语、化学，都能提供专业的错题分析服务。')
    add_formatted_paragraph(doc,
        '（3）实现"以练促学"的闭环——错题录入 → AI分析 → 针对性刷题 → 复习巩固 → 能力提升，形成完整的学习提升路径。')
    add_formatted_paragraph(doc,
        '（4）支持多端同步——Web端、移动端、小程序端数据互通，随时随地管理错题、刷题复习。')

    add_heading_styled(doc, '1.3 初期技术构想', level=2)
    add_formatted_paragraph(doc, '项目初期，团队对技术方案进行了初步构想：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）前端方面：考虑使用React或Vue构建现代化的单页应用（SPA），提供流畅的交互体验。')
    add_formatted_paragraph(doc,
        '（2）后端方面：考虑使用Node.js（Express）或Python（Flask/Django），提供RESTful API服务。')
    add_formatted_paragraph(doc,
        '（3）AI方面：考虑对接大语言模型API（如ChatGPT、DeepSeek等），实现智能分析功能。')
    add_formatted_paragraph(doc,
        '（4）OCR方面：考虑集成Tesseract OCR或调用云端视觉模型API，实现拍照文字识别。')
    add_formatted_paragraph(doc,
        '（5）数据库方面：考虑使用MySQL或PostgreSQL，满足多用户场景的需求。')
    add_formatted_paragraph(doc,
        '在后续的可行性评估和资源约束分析中，团队对技术方案进行了务实调整（见下一章）。',
        first_line_indent=0.74)

    add_heading_styled(doc, '1.4 初期项目目标（MVP定义）', level=2)
    add_formatted_paragraph(doc, '团队在立项阶段就明确了"最小可行产品（MVP）"的范围：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）用户能够注册和登录系统。'
        '（2）用户能够手动录入错题，包括题目文字、选项、答案、解析和知识点标签。'
        '（3）用户能够查看和管理自己的错题本（列表浏览、按学科/标签筛选）。'
        '（4）用户能够基于错题本进行随机刷题练习（出题、答题、批改、成绩）。'
        '（5）系统能够部署到外网，让他人可以通过浏览器访问。')
    add_formatted_paragraph(doc,
        '以上MVP目标已全部实现。AI智能分析、拍照OCR、个性化复习计划等高级功能被纳入后续迭代计划。',
        first_line_indent=0.74)


    # ==== 二、项目可行性评估 ====
    add_heading_styled(doc, '二、项目可行性评估', level=1)

    add_heading_styled(doc, '2.1 技术可行性', level=2)
    add_formatted_paragraph(doc, '经过团队对技术方案的评估，得出以下结论：', first_line_indent=0.74)
    add_heading_styled(doc, '2.1.1 前端技术方案评估', level=3)
    add_table_from_data(doc,
        ['方案', '优点', '缺点', '结论'],
        [
            ['React + Vite', '组件化强、生态完善', '学习成本高、需额外构建工具', '学习周期长，不选'],
            ['Vue 3 + Vite', '上手相对简单', '仍需前后端分离部署', '开发周期较长，不选'],
            ['Jinja2模板+服务端渲染', '零构建工具、部署简单', '交互体验不如SPA', '✅ 最适合课程项目'],
        ])
    add_heading_styled(doc, '2.1.2 后端技术方案评估', level=3)
    add_table_from_data(doc,
        ['方案', '优点', '缺点', '结论'],
        [
            ['Django', '功能全面、自带后台', '框架重量大、配置复杂', '过于庞大，不选'],
            ['Node.js/Express', '异步性能好', '需额外学习Node生态', '学习成本高，不选'],
            ['Flask', '轻量灵活、Python易用', '功能需自行扩展', '✅ 最佳平衡点'],
        ])
    add_heading_styled(doc, '2.1.3 数据库方案评估', level=3)
    add_table_from_data(doc,
        ['方案', '优点', '缺点', '结论'],
        [
            ['MySQL', '功能强大、业界标准', '需独立安装、配置复杂', '超出课程需求'],
            ['PostgreSQL', '功能最强、支持JSON', '同样需独立安装服务', '超出课程需求'],
            ['SQLite', '零配置、文件型数据库', '并发性能有限', '✅ 课程项目最佳选择'],
        ])
    add_heading_styled(doc, '2.1.4 AI方案评估', level=3)
    add_formatted_paragraph(doc,
        '项目初期拟集成AI大模型接口实现智能分析功能。经调研分析：'
        '（1）DeepSeek API——文本分析首选，约1元/百万token，性价比极高；'
        '（2）Qwen-VL——OCR识别首选，中文效果优秀。两大接口均兼容OpenAI API格式，便于统一调用。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '但考虑到课程项目的开发周期、API费用预算以及演示稳定性，团队决策：'
        '当前版本通过手动输入知识点标签的方式替代AI自动分析，AI接口预留设计方案已纳入后续迭代计划。'
        '这一决策确保项目在可控范围内顺利完成MVP交付，同时保持技术方案的延展性。',
        first_line_indent=0.74)

    add_heading_styled(doc, '2.2 经济可行性', level=2)
    add_formatted_paragraph(doc, '本项目经济可行性分析如下：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['成本项', '方案', '月费用', '说明'],
        [
            ['服务器', 'Render免费计划', '¥0', '550小时/月，冷启动稍慢但免费'],
            ['数据库', 'SQLite（本地文件）', '¥0', '无需付费数据库服务'],
            ['域名', 'onrender.com子域名', '¥0', '自动获得免费HTTPS子域名'],
            ['AI API', '手动标签替代AI', '¥0', '当前不调用AI接口'],
            ['开发工具', 'VS Code + Git', '¥0', '免费开源工具'],
            ['AI API（预留）', 'DeepSeek + Qwen-VL', '约¥5/月', '后续可选，按量付费'],
        ])
    add_formatted_paragraph(doc,
        '综上所述，项目在MVP阶段的总成本为零，完全符合课程项目的预算约束。'
        '未来如需集成AI服务，成本也控制在极低水平（约5元/月），具有极高的经济可行性。',
        first_line_indent=0.74)

    add_heading_styled(doc, '2.3 操作可行性', level=2)
    add_formatted_paragraph(doc, '本项目的操作可行性评估如下：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）用户接受度：目标用户（学生）具有明确的错题管理需求，且对智能化学习工具接受度高。'
        '系统设计遵循简洁直观的原则，注册后即可快速上手使用，无需培训。')
    add_formatted_paragraph(doc,
        '（2）技术支持：系统采用Flask + SQLite轻量架构，部署和维护极其简便。'
        '开发团队具备Python编程和Web开发基础，能够自主完成开发、测试和运维工作。')
    add_formatted_paragraph(doc,
        '（3）数据安全：用户密码采用Werkzeug散列加密存储，所有数据通过ORM参数化查询防护SQL注入，'
        '满足基本的系统安全要求。')

    add_heading_styled(doc, '2.4 进度可行性', level=2)
    add_formatted_paragraph(doc, '项目在课程给定的16周时间框架内具有充分的进度可行性：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）技术栈学习成本低：Python + Flask + SQLite的学习曲线平缓，团队快速上手。')
    add_formatted_paragraph(doc,
        '（2）MVP范围可控：初期明确界定了最小可行产品范围，避免范围蔓延（scope creep）。')
    add_formatted_paragraph(doc,
        '（3）分阶段迭代：采用敏捷迭代模型，每个阶段有明确的交付物和检查点。')
    add_formatted_paragraph(doc,
        '（4）时间缓冲充足：计划预留了15%左右的缓冲时间应对意外情况。')

    add_heading_styled(doc, '2.5 支持性信息汇总', level=2)
    add_formatted_paragraph(doc, '综合评估后，项目获得了以下支持性信息的确认：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）课程指导老师对项目选题给予了肯定意见。')
    add_formatted_paragraph(doc,
        '（2）多位同学对项目提供了用户需求反馈。')
    add_formatted_paragraph(doc,
        '（3）技术方案经查阅官方文档和社区资料，验证了技术路线的合理性。')
    add_formatted_paragraph(doc,
        '（4）部署方案经实际测试（本地开发服务器 + Render平台部署演练），确认可正常运行。')


    # ==== 三、项目范围 ====
    add_heading_styled(doc, '三、项目范围', level=1)

    add_heading_styled(doc, '3.1 范围定义', level=2)
    add_formatted_paragraph(doc,
        '项目范围定义了系统应包含和不包含的功能边界，确保团队在规定时间内聚焦核心价值交付。',
        first_line_indent=0.74)

    add_heading_styled(doc, '3.2 范围内功能（In-Scope）', level=2)
    add_formatted_paragraph(doc, '以下功能被明确纳入本次项目的开发范围：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['功能模块', '具体内容', '优先级'],
        [
            ['用户管理', '注册、登录、退出、密码加密、记住我', 'P0'],
            ['错题录入', '手动表单录入（题目、选项、答案、解析、标签、难度）', 'P0'],
            ['错题管理', '列表浏览、学科筛选、标签筛选、标记掌握、删除', 'P0'],
            ['刷题系统', '选择题量/学科/标签筛选、随机出题、答题批改、成绩报告', 'P0'],
            ['仪表盘', '错题统计、掌握率、学科分布、最近错题', 'P1'],
            ['知识图谱', '标签云展示、点击标签跳转筛选', 'P1'],
            ['响应式设计', '桌面端/移动端自适应布局', 'P1'],
            ['外网部署', 'Render平台部署、HTTPS访问', 'P0'],
            ['文档交付', '中期报告、总结报告、设计报告、部署指南', 'P0'],
        ])

    add_heading_styled(doc, '3.3 范围外功能（Out-of-Scope）', level=2)
    add_formatted_paragraph(doc, '以下功能被明确排除在本次项目的开发范围之外，列入后续迭代计划：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['功能', '排除原因', '计划'],
        [
            ['AI大模型接口', 'API成本、开发周期、演示稳定性', '后续版本'],
            ['拍照OCR识别', '视觉模型API成本和集成复杂度', '后续版本'],
            ['间隔重复算法', '课程项目演示的MVP已满足需求', '后续版本'],
            ['微信小程序', '需额外开发小程序版本', '后续考虑'],
            ['第三方登录', '需申请开发者资质', '后续考虑'],
            ['数据导出PDF', 'MVP阶段优先保证核心功能', '后续版本'],
            ['管理员后台', '课程项目规模不需要', '后续考虑'],
            ['手机App', '超出课程项目范围', '远期规划'],
            ['语音输入', '技术方案尚不成熟', '远期规划'],
        ])

    add_heading_styled(doc, '3.4 范围管理策略', level=2)
    add_formatted_paragraph(doc, '为防止范围蔓延，团队采取了以下范围管理策略：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）MVP优先：在项目初期就明确定义了最小可行产品的范围，作为优先级最高的交付目标。')
    add_formatted_paragraph(doc,
        '（2）需求冻结：在完成需求分析后，冻结功能列表，新的需求建议记录在"后续迭代"清单中而非立即实施。')
    add_formatted_paragraph(doc,
        '（3）优先级矩阵：使用P0（必做）、P1（应做）、P2（可做）三级分类，确保核心功能优先完成。')
    add_formatted_paragraph(doc,
        '（4）定期评审：每个开发阶段结束前，团队评审当前完成的功能是否与计划范围一致，及时纠正偏差。')


    # ==== 四、项目概述与成果 ====
    add_heading_styled(doc, '四、项目概述与成果', level=1)

    add_heading_styled(doc, '4.1 项目简介', level=2)
    add_formatted_paragraph(doc,
        '"AI智能错题本与个性化刷题系统"是一个面向学生的智能学习管理Web应用。系统采用Flask + SQLite轻量级技术架构，'
        '实现了错题的手动录入、分类管理、随机刷题、学习数据统计等核心功能。项目以"让每次错误都成为进步的阶梯"为设计理念，'
        '致力于帮助学生高效管理错题、精准定位薄弱点、科学安排复习计划。',
        first_line_indent=0.74)

    add_heading_styled(doc, '4.2 项目完成情况', level=2)
    add_formatted_paragraph(doc,
        '项目按照计划时间线顺利推进，所有核心功能均已完成开发并通过测试。在中期汇报后，根据反馈意见进行了移动端适配优化和部署方案完善。'
        '最终交付物包括：完整的Flask后端代码、11个前端页面、CSS响应式样式、数据库初始化脚本、部署配置文件（Procfile、requirements.txt）'
        '以及三份课程报告文档。',
        first_line_indent=0.74)

    add_heading_styled(doc, '4.3 交付物清单', level=2)
    add_table_from_data(doc,
        ['类别', '交付物', '说明'],
        [
            ['源代码', 'app.py（520行）', 'Flask后端主程序，包含全部路由和数据库模型'],
            ['源代码', 'templates/（11个HTML文件）', 'Jinja2模板，含登录、仪表盘、错题本、刷题、知识图谱'],
            ['静态资源', 'style.css（500行）', '全局响应式CSS样式，支持明暗主题'],
            ['配置', 'requirements.txt', 'Python依赖包列表'],
            ['配置', 'Procfile', '云平台部署启动配置'],
            ['文档', '项目中期进度报告', '需求分析、架构设计、进度总结'],
            ['文档', '项目总结报告', '初期愿景、可行性评估、范围、流程、WBS、会议记录'],
            ['文档', '系统设计报告', '技术架构、数据库设计、接口文档'],
            ['文档', '部署指南', 'Render/ngrok部署步骤'],
        ])


    # ==== 五、系统功能 ====
    add_heading_styled(doc, '五、系统功能', level=1)
    add_heading_styled(doc, '二、系统功能', level=1)

    add_heading_styled(doc, '2.1 功能全景', level=2)
    add_formatted_paragraph(doc, '系统包含六大核心功能模块：', first_line_indent=0.74)

    features = [
        ('认证模块', '用户注册、登录、退出、记住我功能。密码采用Werkzeug散列加密存储。'),
        ('仪表盘模块', '展示错题总数、已掌握数量、掌握率、知识点标签数。按学科展示错题分布。最近错题快捷查看。'),
        ('错题本模块', '手动录入错题（题目、4个选项、正确答案、用户答案、解析、知识点标签、难度）。按学科和标签筛选。标记掌握状态。删除错题。'),
        ('刷题模块', '选择题量（3/5/10/15/20题）和筛选条件。逐题答题，点击选项自动批改。刷题结果报告（得分、每题详情、正确/错误状态）。'),
        ('知识图谱模块', '以标签云展示知识点分布，标签大小反映错题数量。点击标签跳转至对应错题列表。'),
        ('响应式与部署', '移动端适配的汉堡菜单导航。Render平台HTTPS外网部署。gunicorn生产级WSGI服务器。'),
    ]
    for name, desc in features:
        add_heading_styled(doc, name, level=3)
        add_formatted_paragraph(doc, desc, first_line_indent=0.74)

    add_heading_styled(doc, '2.2 用户使用流程', level=2)
    add_formatted_paragraph(doc,
        '注册账号 → 登录系统 → 添加错题（填写题目内容、选项、正确答案、知识点标签等） → '
        '在错题本中查看、管理错题 → 进入刷题页面，选择筛选项和题数 → 逐题作答 → 查看成绩报告 → '
        '在知识图谱中查看知识点分布 → 根据需要继续添加错题或再次刷题',
        first_line_indent=0.74)

    # ==== 六、技术架构 ====
    add_heading_styled(doc, '六、技术架构', level=1)

    add_heading_styled(doc, '3.1 整体架构', level=2)
    add_formatted_paragraph(doc,
        '本系统采用经典的B/S（Browser/Server）三层架构：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）展示层（HTML/CSS）：用户界面由Jinja2模板引擎渲染，使用HTML5语义标签和CSS3响应式布局实现。')
    add_formatted_paragraph(doc,
        '（2）业务逻辑层（Flask）：处理HTTP请求与响应、用户认证与授权、业务数据校验与处理、会话管理。'
        '路由层负责URL映射，模型层（SQLAlchemy ORM）负责数据库交互。')
    add_formatted_paragraph(doc,
        '（3）数据层（SQLite）：存储用户信息、学科数据和错题记录。通过Flask-SQLAlchemy实现ORM映射，'
        '支持自动建表和关系查询。')

    add_heading_styled(doc, '3.2 技术选型理由', level=2)
    add_formatted_paragraph(doc,
        '选择Flask框架的原因：①轻量级——核心仅约10个依赖包，安装和启动极其简便；②灵活——路由设计直观，'
        '模板渲染高效，适合中小型Web应用；③学习曲线平缓——Python语言友好，课程项目团队成员易于上手；'
        '④生态成熟——Flask-Login、Flask-SQLAlchemy等扩展丰富，社区支持完善。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '选择SQLite数据库的原因：①零配置——无需单独安装数据库服务，开箱即用；②嵌入式——数据库文件直接存储在项目目录中，'
        '备份和迁移仅需复制一个文件；③足够用——课程项目单用户、千条数据以内的场景下性能完全满足需求。',
        first_line_indent=0.74)

    add_heading_styled(doc, '3.3 数据库ER图', level=2)
    add_formatted_paragraph(doc, '系统包含三张数据表，关系如下：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        'users（用户表）1 ── N error_questions（错题表） —— 每个用户可以有多道错题\n'
        'subjects（学科表）1 ── N error_questions（错题表） —— 每道错题属于一个学科\n'
        'users（用户表）1 ── 1 profile（用户画像） —— 每个用户有一个扩展信息记录',
        font_size=11)

    # ==== 七、项目开发流程 ====
    add_heading_styled(doc, '七、项目开发流程', level=1)

    add_heading_styled(doc, '7.1 开发生命周期模型', level=2)
    add_formatted_paragraph(doc,
        '本项目采用简化的敏捷迭代开发模型。选择该模型的理由如下：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）项目需求在初期虽已基本明确，但在开发过程中仍可能根据实际体验进行微调，敏捷模型能更好适应需求变化。')
    add_formatted_paragraph(doc,
        '（2）课程项目的开发时间有限（16周），不适合采用重型瀑布模型。敏捷模型的分阶段交付能更早地产生可见成果。')
    add_formatted_paragraph(doc,
        '（3）团队规模小（个人项目），敏捷模型中的Daily Standup、Sprint Planning等仪式可以最大程度简化。')

    add_heading_styled(doc, '7.2 开发流程阶段', level=2)
    add_formatted_paragraph(doc, '整个项目开发流程分为以下阶段：', first_line_indent=0.74)

    add_heading_styled(doc, '第一阶段：启动与规划（第1-2周）', level=3)
    add_formatted_paragraph(doc,
        '· 确定项目选题和核心功能方向\n'
        '· 进行市场调研和可行性分析\n'
        '· 完成技术选型和技术预研\n'
        '· 制定项目开发计划和里程碑\n'
        '· 搭建开发环境（安装Python、Flask、SQLite等）',
        font_size=11)

    add_heading_styled(doc, '第二阶段：需求分析与设计（第3-4周）', level=3)
    add_formatted_paragraph(doc,
        '· 编写用户故事和功能需求文档\n'
        '· 设计系统架构（三层B/S架构）\n'
        '· 设计数据库ER图和表结构\n'
        '· 设计页面原型和交互流程\n'
        '· 编写接口定义文档',
        font_size=11)

    add_heading_styled(doc, '第三阶段：迭代开发（第5-12周）', level=3)
    add_formatted_paragraph(doc,
        '· Sprint 1（第5-6周）：用户认证系统（注册、登录、退出、记住我）\n'
        '· Sprint 2（第7-8周）：错题本核心功能（录入、列表、筛选、详情、标记、删除）\n'
        '· Sprint 3（第9-10周）：刷题系统（出题配置、答题交互、自动批改、成绩报告）\n'
        '· Sprint 4（第11-12周）：仪表盘与知识图谱（数据统计、标签云、前端优化）',
        font_size=11)

    add_heading_styled(doc, '第四阶段：测试与部署（第13-14周）', level=3)
    add_formatted_paragraph(doc,
        '· 功能集成测试和Bug修复\n'
        '· 移动端兼容性测试\n'
        '· 部署到Render云平台\n'
        '· 外网访问验证\n'
        '· 中期进度报告撰写',
        font_size=11)

    add_heading_styled(doc, '第五阶段：收尾与交付（第15-16周）', level=3)
    add_formatted_paragraph(doc,
        '· 代码整理与注释完善\n'
        '· 部署文档编写\n'
        '· 项目总结报告撰写\n'
        '· 系统设计报告撰写\n'
        '· 最终成果打包交付',
        font_size=11)

    add_heading_styled(doc, '7.3 版本管理策略', level=2)
    add_formatted_paragraph(doc,
        '项目采用Git进行版本管理。虽然项目为个人开发，但仍严格遵循了以下版本管理规范：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）主分支（main）：始终保持可运行状态，每个Sprint结束后的稳定版本合并至main。')
    add_formatted_paragraph(doc,
        '（2）功能分支（feature/*）：每开发一个新功能模块时，从main创建feature分支，开发完成后合并回去。如feature/auth、feature/error-book、feature/practice等。')
    add_formatted_paragraph(doc,
        '（3）提交规范：使用清晰的commit message，格式为"[模块] 简要描述"。如"[Auth] 实现用户注册功能"、"[ErrorBook] 添加错题列表学科筛选"。')


    # ==== 八、项目WBS工作分解 ====
    add_heading_styled(doc, '八、项目WBS工作分解结构', level=1)

    add_heading_styled(doc, '8.1 WBS概述', level=2)
    add_formatted_paragraph(doc,
        '工作分解结构（Work Breakdown Structure, WBS）是将项目可交付成果和项目工作分解为较小的、更易于管理的组分的过程。'
        '以下是本项目的四级WBS分解：',
        first_line_indent=0.74)

    add_heading_styled(doc, '8.2 WBS结构表', level=2)
    add_table_from_data(doc,
        ['WBS编号', '工作包名称', '所属阶段', '预估工时', '前置依赖'],
        [
            ['1.0', 'AI智能错题本与个性化刷题系统', '全部', '16周', '无'],
            ['1.1', '项目启动与规划', '阶段一', '2周', '无'],
            ['1.1.1', '选题与创意构思', '阶段一', '0.5周', '无'],
            ['1.1.2', '市场调研与可行性分析', '阶段一', '0.5周', '1.1.1'],
            ['1.1.3', '技术选型与预研', '阶段一', '0.5周', '1.1.2'],
            ['1.1.4', '开发计划与里程碑制定', '阶段一', '0.5周', '1.1.3'],
            ['1.2', '需求分析与设计', '阶段二', '2周', '1.1'],
            ['1.2.1', '用户故事与功能需求编写', '阶段二', '0.5周', '1.1'],
            ['1.2.2', '系统架构设计', '阶段二', '0.5周', '1.2.1'],
            ['1.2.3', '数据库ER图与表结构设计', '阶段二', '0.5周', '1.2.2'],
            ['1.2.4', '页面原型与交互流程设计', '阶段二', '0.5周', '1.2.2'],
            ['1.3', '迭代开发', '阶段三', '8周', '1.2'],
            ['1.3.1', 'Sprint 1：用户认证系统', '阶段三', '2周', '1.2'],
            ['1.3.1.1', '数据库模型定义（User表）', '阶段三', '0.5周', '1.2.3'],
            ['1.3.1.2', '注册功能（含表单校验）', '阶段三', '0.5周', '1.3.1.1'],
            ['1.3.1.3', '登录与退出功能', '阶段三', '0.5周', '1.3.1.2'],
            ['1.3.1.4', '记住我与会话保持', '阶段三', '0.5周', '1.3.1.3'],
            ['1.3.2', 'Sprint 2：错题本功能', '阶段三', '2周', '1.3.1'],
            ['1.3.2.1', '数据库模型定义（ErrorQuestion表）', '阶段三', '0.5周', '1.3.1'],
            ['1.3.2.2', '错题录入页面与表单', '阶段三', '0.5周', '1.3.2.1'],
            ['1.3.2.3', '错题列表与筛选功能', '阶段三', '0.5周', '1.3.2.2'],
            ['1.3.2.4', '错题详情与标记掌握', '阶段三', '0.5周', '1.3.2.3'],
            ['1.3.3', 'Sprint 3：刷题系统', '阶段三', '2周', '1.3.2'],
            ['1.3.3.1', '刷题配置页面与出题逻辑', '阶段三', '0.5周', '1.3.2'],
            ['1.3.3.2', '答题页面与选项交互', '阶段三', '0.5周', '1.3.3.1'],
            ['1.3.3.3', '自动批改与结果反馈', '阶段三', '0.5周', '1.3.3.2'],
            ['1.3.3.4', '成绩报告页面', '阶段三', '0.5周', '1.3.3.3'],
            ['1.3.4', 'Sprint 4：仪表盘与附加功能', '阶段三', '2周', '1.3.3'],
            ['1.3.4.1', '仪表盘数据统计', '阶段三', '0.5周', '1.3.3'],
            ['1.3.4.2', '知识图谱（标签云）', '阶段三', '0.5周', '1.3.4.1'],
            ['1.3.4.3', '响应式前端优化', '阶段三', '0.5周', '1.3.4.2'],
            ['1.3.4.4', 'CSS动画与交互细节', '阶段三', '0.5周', '1.3.4.3'],
            ['1.4', '测试与部署', '阶段四', '2周', '1.3'],
            ['1.4.1', '功能集成测试', '阶段四', '0.5周', '1.3'],
            ['1.4.2', '移动端兼容性测试', '阶段四', '0.5周', '1.4.1'],
            ['1.4.3', 'Render云平台部署', '阶段四', '0.5周', '1.4.2'],
            ['1.4.4', '外网访问验证与整改', '阶段四', '0.5周', '1.4.3'],
            ['1.5', '文档与收尾', '阶段五', '2周', '1.4'],
            ['1.5.1', '代码整理与注释完善', '阶段五', '0.5周', '1.4'],
            ['1.5.2', '部署指南与使用说明', '阶段五', '0.5周', '1.5.1'],
            ['1.5.3', '中期进度报告', '阶段五', '0.5周', '1.4'],
            ['1.5.4', '项目总结报告', '阶段五', '0.5周', '1.5.3'],
            ['1.5.5', '系统设计报告', '阶段五', '0.5周', '1.5.4'],
            ['1.5.6', '最终成果打包', '阶段五', '0.5周', '1.5.5'],
        ])


    # ==== 九、项目会议记录 ====
    add_heading_styled(doc, '九、项目会议记录', level=1)

    add_heading_styled(doc, '9.1 会议记录说明', level=2)
    add_formatted_paragraph(doc,
        '项目开发过程中，团队（个人项目）定期进行阶段性自我审查，并在关键节点与课程指导老师进行进度沟通。'
        '以下是各关键节点的会议记录摘要：',
        first_line_indent=0.74)

    add_heading_styled(doc, '9.2 会议记录清单', level=2)

    add_heading_styled(doc, '9.2.1 项目启动会议', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年3月1日'],
            ['会议主题', '项目选题确认与方向讨论'],
            ['参与人员', '薛文锐（项目负责人）、课程指导老师'],
            ['会议内容', '1. 讨论并确定项目选题为"AI智能错题本与个性化刷题系统"；\n'
                        '2. 明确目标用户为中学生和大学生；\n'
                        '3. 初步确定核心功能范围（错题录入、AI分析、刷题系统）；\n'
                        '4. 讨论技术方案的初步设想。'],
            ['会议结论', '选题通过。建议：优先完成基础版本，AI功能作为后续迭代目标。'],
            ['后续行动', '开展详细的需求分析和技术预研。'],
        ])

    add_heading_styled(doc, '9.2.2 需求分析评审会', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年3月15日'],
            ['会议主题', '需求分析文档评审'],
            ['参与人员', '薛文锐'],
            ['会议内容', '1. 确认了9项核心功能需求（FR-01至FR-09）；\n'
                        '2. 定义了6项非功能性需求（NFR-01至NFR-06）；\n'
                        '3. 讨论了AI功能在当前版本中的替代方案（手动标签）。'],
            ['会议结论', '需求范围合理，明确将AI功能推迟至后续迭代。'],
            ['后续行动', '开始数据库设计和系统架构设计。'],
        ])

    add_heading_styled(doc, '9.2.3 技术选型评审会', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年3月28日'],
            ['会议主题', '技术选型最终确认'],
            ['参与人员', '薛文锐'],
            ['会议内容', '1. 对比评估了React/Vue（前后端分离）与Flask+Jinja2（服务端渲染）两种方案；\n'
                        '2. 对比评估了MySQL/PostgreSQL与SQLite的数据库方案；\n'
                        '3. 确定最终技术栈：Flask + SQLite + Jinja2 + gunicorn；\n'
                        '4. 确定部署平台：Render（免费）或ngrok（内网穿透）作为备选。'],
            ['会议结论', '采用Flask + SQLite轻量级技术栈，优先Render部署方案。'],
            ['后续行动', '搭建开发环境，开始编码。'],
        ])

    add_heading_styled(doc, '9.2.4 Sprint 1 回顾会', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年4月20日'],
            ['会议主题', '认证系统开发完成回顾'],
            ['参与人员', '薛文锐'],
            ['会议内容', '1. 演示用户注册、登录、退出功能；\n'
                        '2. 问题反馈：重启Flask后session丢失导致用户被强制退出；\n'
                        '3. 定位问题：SECRET_KEY使用os.urandom()每次随机生成。'],
            ['会议结论', '将SECRET_KEY改为固定值以解决session持久化问题。'],
            ['后续行动', '开始错题本功能的开发。'],
        ])

    add_heading_styled(doc, '9.2.5 中期进度审查会', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年5月15日'],
            ['会议主题', '中期进度汇报与演示'],
            ['参与人员', '薛文锐、课程指导老师'],
            ['会议内容', '1. 演示了当前已完成的所有功能（认证、错题本、刷题、仪表盘）；\n'
                        '2. 汇报了项目进度（按计划推进，前四个阶段已完成）；\n'
                        '3. 讨论了部署方案和后续工作计划；\n'
                        '4. 指导老师对移动端体验提出了优化建议。'],
            ['会议结论', '项目进度良好，建议加强移动端适配并尽快完成外网部署。'],
            ['后续行动', '1. 移动端响应式优化；2. 部署到Render平台；3. 撰写中期报告。'],
        ])

    add_heading_styled(doc, '9.2.6 部署上线评审会', level=3)
    add_table_from_data(doc,
        ['项目', '内容'],
        [
            ['会议日期', '2026年6月1日'],
            ['会议主题', '外网部署验证与最终功能确认'],
            ['参与人员', '薛文锐'],
            ['会议内容', '1. 确认Render平台部署成功，外网可正常访问；\n'
                        '2. 逐项测试了所有功能的完整性；\n'
                        '3. 确认代码已推送至GitHub仓库；\n'
                        '4. 确认三份报告文档的撰写范围和大纲。'],
            ['会议结论', '系统部署成功，功能完整，进入文档交付阶段。'],
            ['后续行动', '撰写项目总结报告、系统设计报告。'],
        ])

    add_heading_styled(doc, '9.3 项目沟通总结', level=2)
    add_formatted_paragraph(doc, '项目周期内，团队通过多种渠道保持了有效的沟通：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）正式会议：共召开6次正式评审会议（如上所列），覆盖了从立项到收尾的完整生命周期。')
    add_formatted_paragraph(doc,
        '（2）日常沟通：通过即时通讯工具进行日常问题讨论和技术交流。')
    add_formatted_paragraph(doc,
        '（3）文档沟通：通过共享文档和GitHub代码仓库保持信息的透明性和可追溯性。')
    add_formatted_paragraph(doc,
        '（4）导师沟通：在项目关键节点与课程指导老师进行进度沟通，获取指导和建议。')


    # ==== 十、成果展示 ====
    add_heading_styled(doc, '十、成果展示', level=1)

    add_heading_styled(doc, '5.1 系统功能清单', level=2)
    add_table_from_data(doc,
        ['功能', '状态', '路由'],
        [
            ['用户注册', '完成', '/register'],
            ['用户登录', '完成', '/login'],
            ['用户退出', '完成', '/logout'],
            ['学习仪表盘', '完成', '/dashboard'],
            ['错题本列表', '完成', '/error-book'],
            ['添加错题', '完成', '/error-book/add'],
            ['错题详情', '完成', '/error-book/<id>'],
            ['标记掌握', '完成', '/error-book/<id>/toggle-mastered'],
            ['删除错题', '完成', '/error-book/<id>/delete'],
            ['刷题入口', '完成', '/practice'],
            ['开始刷题', '完成', '/practice/start'],
            ['答题页面', '完成', '/practice/session'],
            ['提交答案', '完成', '/practice/answer'],
            ['下一题', '完成', '/practice/next'],
            ['刷题结果', '完成', '/practice/result'],
            ['知识图谱', '完成', '/knowledge-map'],
        ])

    add_heading_styled(doc, '5.2 代码统计', level=2)
    add_table_from_data(doc,
        ['指标', '数值'],
        [
            ['Python代码行数', '约520行'],
            ['HTML模板行数', '约800行'],
            ['CSS样式行数', '约500行'],
            ['总代码行数', '约1820行'],
            ['数据库表数', '3张'],
            ['API路由数', '16个'],
            ['前端页面数', '11个'],
        ])

    # ==== 十一、收获与反思 ====
    add_heading_styled(doc, '十一、收获与反思', level=1)

    add_heading_styled(doc, '11.1 技术收获', level=2)
    gains = [
        '（1）深入理解了Flask Web框架的工作原理，包括路由分发、请求上下文、模板渲染等核心机制。',
        '（2）掌握了SQLAlchemy ORM的使用方法，能够熟练地进行数据库建模、关系查询和数据迁移。',
        '（3）学习了Flask-Login的认证流程，理解了基于session的用户状态管理机制。',
        '（4）实践了响应式Web设计，使用CSS媒体查询和弹性布局实现了移动端适配。',
        '（5）体验了从本地开发到云端部署的完整链路，理解了WSGI服务器（gunicorn）在生产环境中的作用。',
    ]
    for g in gains:
        add_formatted_paragraph(doc, g, first_line_indent=0.74)

    add_heading_styled(doc, '11.2 项目管理收获', level=2)
    mgmt_gains = [
        '（1）学会了制定合理的项目计划，将大目标拆解为可执行的小任务，按优先级有序推进。',
        '（2）体会到"先做核心功能"的重要性——在有限的课程时间内，优先保证核心用户流程（注册→录入→刷题）的完整性。',
        '（3）认识到文档编写与代码开发同等重要——及时的文档记录有助于后续的维护和改进。',
        '（4）理解了持续集成和版本管理（Git）在项目开发中的价值。',
    ]
    for g in mgmt_gains:
        add_formatted_paragraph(doc, g, first_line_indent=0.74)

    add_heading_styled(doc, '11.3 不足之处', level=2)
    add_formatted_paragraph(doc,
        '（1）AI功能未实际集成：受限于API成本和开发时间，AI错题分析功能目前通过知识点标签手动替代，后续仍需对接AI服务。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（2）单用户数据隔离不足：当前版本数据库设计支持多用户，但在刷题出题逻辑中未严格按用户过滤，多用户场景需进一步完善。',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（3）自动化测试欠缺：由于时间限制，系统未编写单元测试和集成测试，功能验证主要依赖手动测试。',
        first_line_indent=0.74)

    # ==== 十二、改进方向 ====
    add_heading_styled(doc, '十二、改进方向', level=1)
    improvements = [
        '（1）AI集成：对接大语言模型API（如DeepSeek），实现错题自动分析、知识点自动提取和复习计划生成。',
        '（2）数据库升级：从SQLite迁移至PostgreSQL，使用Supabase免费PostgreSQL服务，提升并发能力和数据持久性。',
        '（3）前后端分离：将前端重构为React/Vue单页应用，后端改为纯RESTful API，提升用户体验和可维护性。',
        '（4）拍照OCR：集成视觉AI模型实现拍照录入错题，大幅降低录入门槛。',
        '（5）间隔重复算法：引入SM-2算法，根据答题表现动态调整错题的复习时间间隔，提升复习效率。',
        '（6）数据导出：支持将错题本导出为PDF或打印格式，方便线下复习使用。',
        '（7）测试覆盖：编写单元测试和端到端测试，确保代码质量和功能稳定性。',
    ]
    for imp in improvements:
        add_formatted_paragraph(doc, imp, first_line_indent=0.74)

    # 保存
    filepath = os.path.join(OUTPUT_DIR, '2-项目总结报告.docx')
    doc.save(filepath)
    print(f'已生成：{filepath}')
    return filepath


# ======================================================================
#                          报告三：系统设计报告
# ======================================================================
def generate_design_report():
    doc = setup_document()
    add_title_page(doc, 'AI 智能错题本与个性化刷题系统', '系统设计报告')

    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、系统概述 ............................................. 1',
        '二、需求分析 ............................................. 2',
        '三、系统架构设计 ........................................ 3',
        '四、数据库设计 ........................................... 5',
        '五、接口设计 ............................................. 6',
        '六、前端页面设计 ........................................ 8',
        '七、安全设计 ............................................. 9',
    ]
    for item in toc_items:
        add_formatted_paragraph(doc, item, font_size=12, space_after=4)
    doc.add_page_break()

    # ==== 一、系统概述 ====
    add_heading_styled(doc, '一、系统概述', level=1)

    add_heading_styled(doc, '1.1 系统定位', level=2)
    add_formatted_paragraph(doc,
        '"AI智能错题本与个性化刷题系统"是一个面向学生的智能学习管理Web应用程序。'
        '系统旨在帮助学生高效管理错题、精准定位知识薄弱点，通过智能化的刷题和复习策略提升学习效果。',
        first_line_indent=0.74)

    add_heading_styled(doc, '1.2 设计目标', level=2)
    add_formatted_paragraph(doc, '本系统的设计目标包括：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）简单易用：界面简洁直观，操作流程清晰，降低学习成本。')
    add_formatted_paragraph(doc,
        '（2）高效管理：支持快速的错题录入、分类筛选和状态追踪。')
    add_formatted_paragraph(doc,
        '（3）智能刷题：基于用户错题数据，智能生成个性化练习题目。')
    add_formatted_paragraph(doc,
        '（4）轻量部署：采用轻量级技术栈，支持快速部署至免费云平台。')

    # ==== 二、需求分析 ====
    add_heading_styled(doc, '二、需求分析', level=1)

    add_heading_styled(doc, '2.1 功能性需求', level=2)
    add_table_from_data(doc,
        ['编号', '功能名称', '描述', '优先级'],
        [
            ['FR-01', '用户注册', '新用户通过用户名和密码创建账户', '高'],
            ['FR-02', '用户登录', '已注册用户通过用户名和密码登录系统', '高'],
            ['FR-03', '错题录入', '用户手动录入错题，包括题目、选项、答案、知识点标签', '高'],
            ['FR-04', '错题浏览', '查看所有错题列表，支持按学科和标签筛选', '高'],
            ['FR-05', '错题详情', '查看错题的完整信息，包括题目、选项、答案对比、解析', '高'],
            ['FR-06', '错题管理', '标记错题为已掌握、删除错题', '中'],
            ['FR-07', '刷题练习', '从错题本中随机抽取题目生成练习', '高'],
            ['FR-08', '答题批改', '在线作答并自动批改，显示正确答案和解析', '高'],
            ['FR-09', '成绩报告', '刷题结束后展示得分、正确率、每题详情', '中'],
            ['FR-10', '学习统计', '在仪表盘展示错题总数、掌握率、学科分布', '中'],
            ['FR-11', '知识图谱', '以标签云展示知识点分布', '低'],
        ])

    add_heading_styled(doc, '2.2 非功能性需求', level=2)
    add_table_from_data(doc,
        ['编号', '需求类型', '描述'],
        [
            ['NFR-01', '可用性', '界面直观，无需培训即可上手使用'],
            ['NFR-02', '响应性', '页面加载时间不超过3秒'],
            ['NFR-03', '兼容性', '支持主流浏览器（Chrome、Edge、Firefox）'],
            ['NFR-04', '移动适配', '在手机和平板设备上的显示正常'],
            ['NFR-05', '安全性', '用户密码加密存储，防止SQL注入'],
            ['NFR-06', '可部署性', '支持一键部署至免费云平台'],
        ])

    # ==== 三、系统架构设计 ====
    add_heading_styled(doc, '三、系统架构设计', level=1)

    add_heading_styled(doc, '3.1 整体架构', level=2)
    add_formatted_paragraph(doc,
        '系统采用经典的B/S三层架构，分为展示层、业务逻辑层和数据层：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '┌─────────────────────────────────────────┐\n'
        '│           展示层（Presentation）          │\n'
        '│  HTML5 + CSS3 + Jinja2 Template Engine   │\n'
        '│  · 登录/注册 · 仪表盘 · 错题本 · 刷题 · 知识图谱  │\n'
        '├─────────────────────────────────────────┤\n'
        '│          业务逻辑层（Business Logic）      │\n'
        '│  Flask 3.x + Python 3                    │\n'
        '│  · 路由处理 · 认证授权 · 数据校验 · 会话管理   │\n'
        '│  · Flask-Login · Flask-SQLAlchemy        │\n'
        '├─────────────────────────────────────────┤\n'
        '│            数据层（Data）                  │\n'
        '│  SQLite 3（通过SQLAlchemy ORM访问）         │\n'
        '│  · users · subjects · error_questions     │\n'
        '└─────────────────────────────────────────┘',
        font_size=10)

    add_heading_styled(doc, '3.2 技术栈', level=2)
    add_table_from_data(doc,
        ['技术组件', '选型', '版本', '作用'],
        [
            ['编程语言', 'Python', '3.9+', '后端业务逻辑开发'],
            ['Web框架', 'Flask', '3.1.3', 'HTTP路由、请求响应处理'],
            ['ORM', 'Flask-SQLAlchemy', '3.1.1', '数据库对象映射'],
            ['认证', 'Flask-Login', '0.6.3', '用户会话管理'],
            ['密码加密', 'Werkzeug', '3.1.8', '密码哈希与验证'],
            ['数据库', 'SQLite', '3', '数据持久化存储'],
            ['模板引擎', 'Jinja2', '3.1.6', 'HTML页面渲染'],
            ['WSGI Server', 'gunicorn', '23.0.0', '生产环境Web服务'],
            ['云平台', 'Render', '-', '外网部署与HTTPS服务'],
        ])

    add_heading_styled(doc, '3.3 数据流设计', level=2)
    add_formatted_paragraph(doc,
        '典型的数据流程如下：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '（1）用户请求 → Flask路由接收HTTP请求 → 解析参数/表单数据\n'
        '（2）业务处理 → 调用SQLAlchemy查询/更新数据库 → 获取结果\n'
        '（3）模板渲染 → Jinja2将数据注入HTML模板 → 生成完整页面\n'
        '（4）HTTP响应 → 浏览器接收HTML/CSS/JS → 渲染用户界面')

    # ==== 四、数据库设计 ====
    add_heading_styled(doc, '四、数据库设计', level=1)

    add_heading_styled(doc, '4.1 ER图', level=2)
    add_formatted_paragraph(doc, '系统包含三张核心数据表，实体关系如下：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '┌──────────┐        ┌───────────────────┐        ┌──────────┐\n'
        '│  users   │ 1 ─── N │ error_questions    │ N ─── 1 │ subjects │\n'
        '│──────────│        │───────────────────│        │──────────│\n'
        '│ id (PK)  │        │ id (PK)            │        │ id (PK)  │\n'
        '│ username │        │ user_id (FK)       │        │ name     │\n'
        '│ password │        │ subject_id (FK)     │        │ icon     │\n'
        '│ created  │        │ question_text      │        │          │\n'
        '└──────────┘        │ option_a/b/c/d     │        └──────────┘\n'
        '                    │ correct_answer     │\n'
        '                    │ user_answer        │\n'
        '                    │ explanation        │\n'
        '                    │ knowledge_tags     │\n'
        '                    │ difficulty         │\n'
        '                    │ mastered           │\n'
        '                    │ created_at         │\n'
        '                    └───────────────────┘',
        font_size=10)

    add_heading_styled(doc, '4.2 数据字典', level=2)

    add_heading_styled(doc, 'error_questions表（错题表）', level=3)
    add_table_from_data(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PRIMARY KEY', '自增主键'],
            ['user_id', 'INTEGER', 'FOREIGN KEY → users.id', '所属用户'],
            ['subject_id', 'INTEGER', 'FOREIGN KEY → subjects.id', '所属学科'],
            ['question_text', 'TEXT', 'NOT NULL', '题目内容'],
            ['option_a', 'VARCHAR(500)', 'NULL', '选项A'],
            ['option_b', 'VARCHAR(500)', 'NULL', '选项B'],
            ['option_c', 'VARCHAR(500)', 'NULL', '选项C'],
            ['option_d', 'VARCHAR(500)', 'NULL', '选项D'],
            ['correct_answer', 'VARCHAR(10)', 'NOT NULL', '正确答案（A/B/C/D）'],
            ['user_answer', 'VARCHAR(10)', 'NULL', '用户作答'],
            ['explanation', 'TEXT', 'NULL', '题目解析'],
            ['knowledge_tags', 'VARCHAR(500)', 'NULL', '知识点标签（逗号分隔）'],
            ['difficulty', 'INTEGER', 'DEFAULT 3', '难度（1-5）'],
            ['error_count', 'INTEGER', 'DEFAULT 1', '错误次数'],
            ['mastered', 'BOOLEAN', 'DEFAULT FALSE', '是否已掌握'],
            ['created_at', 'DATETIME', 'DEFAULT NOW', '创建时间'],
            ['updated_at', 'DATETIME', 'DEFAULT NOW', '更新时间'],
        ])

    # ==== 五、接口设计 ====
    add_heading_styled(doc, '五、接口设计', level=1)

    add_heading_styled(doc, '5.1 后端路由列表', level=2)
    add_table_from_data(doc,
        ['方法', '路由', '功能', '认证'],
        [
            ['GET/POST', '/login', '登录页面', '否'],
            ['GET/POST', '/register', '注册页面', '否'],
            ['GET', '/logout', '退出登录', '是'],
            ['GET', '/dashboard', '学习仪表盘', '是'],
            ['GET', '/error-book', '错题本列表', '是'],
            ['GET/POST', '/error-book/add', '添加错题', '是'],
            ['GET', '/error-book/<id>', '错题详情', '是'],
            ['POST', '/error-book/<id>/toggle-mastered', '切换掌握状态', '是'],
            ['POST', '/error-book/<id>/delete', '删除错题', '是'],
            ['GET', '/practice', '刷题入口', '是'],
            ['POST', '/practice/start', '开始刷题', '是'],
            ['GET', '/practice/session', '答题页面', '是'],
            ['POST', '/practice/answer', '提交答案（AJAX）', '是'],
            ['POST', '/practice/next', '下一题（AJAX）', '是'],
            ['GET', '/practice/result', '刷题结果', '是'],
            ['GET', '/knowledge-map', '知识图谱', '是'],
        ])

    add_heading_styled(doc, '5.2 关键接口说明', level=2)

    add_heading_styled(doc, '5.2.1 添加错题 POST /error-book/add', level=3)
    add_formatted_paragraph(doc, '请求参数（表单提交）：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['参数', '类型', '必填', '说明'],
        [
            ['subject_id', 'int', '是', '学科ID'],
            ['question_text', 'text', '是', '题目内容'],
            ['option_a~d', 'text', '否', '选项A到D'],
            ['correct_answer', 'text', '是', '正确答案（A/B/C/D）'],
            ['user_answer', 'text', '否', '用户答案'],
            ['explanation', 'text', '否', '题目解析'],
            ['knowledge_tags', 'text', '否', '知识点标签（逗号分隔）'],
            ['difficulty', 'int', '否', '难度（1-5，默认3）'],
        ])

    add_heading_styled(doc, '5.2.2 提交答案 POST /practice/answer（AJAX）', level=3)
    add_formatted_paragraph(doc, '请求参数（表单）：', first_line_indent=0.74)
    add_table_from_data(doc,
        ['参数', '类型', '必填', '说明'],
        [
            ['answer', 'text', '是', '用户选择的选项（A/B/C/D）'],
        ])
    add_formatted_paragraph(doc, '响应示例（JSON）：', first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '{\n'
        '  "success": true,\n'
        '  "user_answer": "A",\n'
        '  "correct_answer": "B",\n'
        '  "is_correct": false\n'
        '}',
        font_size=10)

    # ==== 六、前端页面设计 ====
    add_heading_styled(doc, '六、前端页面设计', level=1)

    add_heading_styled(doc, '6.1 页面结构', level=2)
    add_formatted_paragraph(doc,
        '系统采用服务端渲染（SSR）模式，由Jinja2模板引擎在服务端生成完整的HTML页面后返回给浏览器。'
        '前端页面共11个，分为认证页面组（login、register）、主功能页面组（dashboard、error_book、'
        'add_error、error_detail、practice、practice_session、practice_result、knowledge_map）和公共组件（navbar）。',
        first_line_indent=0.74)

    add_heading_styled(doc, '6.2 响应式设计', level=2)
    add_formatted_paragraph(doc,
        '系统采用CSS媒体查询实现响应式布局：',
        first_line_indent=0.74)
    add_formatted_paragraph(doc,
        '· 桌面端（>768px）：左侧固定侧边栏导航，内容区居中，最大宽度960px')
    add_formatted_paragraph(doc,
        '· 移动端（≤768px）：顶部汉堡菜单替代侧边栏，内容区全宽显示，表格改为纵向堆叠')
    add_formatted_paragraph(doc,
        '· 通用适配：表单输入框、按钮、卡片均使用相对单位和弹性盒模型，自动适配不同屏幕尺寸')

    add_heading_styled(doc, '6.3 配色方案', level=2)
    add_table_from_data(doc,
        ['用途', '颜色', '说明'],
        [
            ['主色调', '#6366f1（靛蓝紫）', '按钮、链接、高亮、进度条'],
            ['成功色', '#22c55e（绿色）', '正确标记、已掌握标记'],
            ['警告色', '#f59e0b（琥珀色）', '中等状态提示'],
            ['错误色', '#ef4444（红色）', '错误标记、删除按钮'],
            ['背景色', '#f5f5f5（浅灰）', '页面背景'],
            ['卡片色', '#ffffff（白色）', '内容卡片背景'],
            ['文字色', '#1f2937（深灰）', '正文文字'],
            ['辅助文字', '#6b7280（中灰）', '描述性文字'],
        ])

    # ==== 七、安全设计 ====
    add_heading_styled(doc, '七、安全设计', level=1)

    add_heading_styled(doc, '7.1 身份认证', level=2)
    add_formatted_paragraph(doc,
        '系统采用基于Session的用户认证机制。用户登录时，Flask-Login在服务端创建加密的Session Cookie，'
        '之后的每次请求通过该Cookie识别用户身份。密码在存储前使用Werkzeug的'
        'generate_password_hash()进行PBKDF2散列处理，无法逆向还原为明文。',
        first_line_indent=0.74)

    add_heading_styled(doc, '7.2 访问控制', level=2)
    add_formatted_paragraph(doc,
        '系统使用Flask-Login的@login_required装饰器保护所有功能页面。未登录用户访问受保护页面时，'
        '自动重定向至登录页面。每个用户可以访问和操作自己的错题数据，通过SQLAlchemy的filter_by(user_id=current_user.id)实现数据隔离。',
        first_line_indent=0.74)

    add_heading_styled(doc, '7.3 输入安全', level=2)
    add_formatted_paragraph(doc,
        '（1）SQL注入防护：使用SQLAlchemy ORM的参数化查询，所有SQL语句均通过参数绑定，杜绝SQL注入风险。'
        '（2）XSS防护：Jinja2模板引擎默认对输出进行HTML转义，防止跨站脚本攻击。'
        '（3）CSRF防护：Flask-SQLAlchemy的Session机制提供了基础的CSRF防御。后续可引入Flask-WTF的CSRF Token增强保护。',
        first_line_indent=0.74)

    # 保存
    filepath = os.path.join(OUTPUT_DIR, '3-系统设计报告.docx')
    doc.save(filepath)
    print(f'已生成：{filepath}')
    return filepath


# ======================================================================
#                              主入口
# ======================================================================
if __name__ == '__main__':
    print("开始生成三份报告...")
    generate_mid_term_report()
    generate_final_report()
    generate_design_report()
    print(f"\n三份报告已生成至：{OUTPUT_DIR}")
    print("文件列表：")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        print(f"  - {f}")
